#选取的地址http://www.81.cn/
import requests
from bs4 import BeautifulSoup
import json
from docx import Document
from docx.shared import Pt

def change_docx(manuscriptDatas):
    i = 1
    for meta_data in manuscriptDatas:
        html_url = meta_data.get('url')
        response = requests.get(html_url)
        response.encoding = 'utf-8'
        html_data = response.text
        # html_data是包含HTML内容的变量
        soup = BeautifulSoup(html_data, 'html.parser')
        # 查找所有带有<p>和<title>标签的内容
        span_tags = soup.find_all(['title','p'])
        text_list = [tag.get_text() for tag in span_tags[:-1]]
        if '看军事 上军号' in text_list and '军事资讯 早知道' in text_list and '扫描下方二维码' in text_list and '下载中国军号客户端' in text_list:
            text_list =  [tag.get_text() for tag in span_tags[:-6]]
        else:
            text_list = [tag.get_text() for tag in span_tags[:-1]]
        print(text_list)
        #创建保存路径
        doc = Document()
        # 设置标题
        doc.add_heading(text_list[0], level=1)
        # 逐个添加文本内容
        for text in text_list[1:]:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            # 设置字体样式
            font = run.font
            font.name = 'Arial'  # 设置字体为Arial
            font.size = Pt(12)  # 设置字体大小为12磅
            run.bold = False
        # 保存文档
        doc.save('D:/study/python_project/KG_IR_gongwen/resourcecode/pachong_information/original_data/'+'报告'+str(i)+'.docx')
        text_list = []
        i = i + 1
def main():
    # 设定你要查询的网站URL
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
        'Origin': 'http://www.mod.gov.cn',
        'Referer': 'http://www.mod.gov.cn/',
        'Accept-Language': 'zh-CN,zh;q=0.9',
    }
    url = 'http://mod-search.mod.gov.cn/api-surface/es/docSearchEasy?pageNumber=1&pageSize=10&returnFields&indexNames=manuscript&highlightType=2&title=%E5%86%9B%E9%98%9F%E6%8A%A5%E5%91%8A&searchType=1&accessToken=eXAiOiJKV1QiLJ9Jpc3MiOiJodHRwOlwvXC9n&channelId=718'
    response = requests.get(url, headers=headers)
    # 检查请求是否成功
    if response.status_code == 200:
        manuscriptDatas = []
       # 将字符串转换为JSON对象
        json_response = json.loads(response.text)
        # 提取data内的数据
        data_urls = json_response.get('data', {}).get('dataList',{})
        for data_url in data_urls:
            data_url = data_url.get('manuscriptData',{})
            manuscriptDatas.append(data_url)
        return manuscriptDatas
    else:
        print('请求失败，状态码：', response.status_code)

if __name__ == "__main__":
    manuscriptDatas = main()
    #将获取的URL进行文本解析，并保存到pdf文件中
    print(manuscriptDatas)
    text_to_docx = change_docx(manuscriptDatas)

